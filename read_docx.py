# -*- coding: utf-8 -*-
"""
读取Word文档内容的Python脚本
"""

try:
    import docx
    
    def read_word_document(file_path):
        """读取Word文档内容"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            # 读取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text.strip())
            
            # 读取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))
            
            return content
            
        except Exception as e:
            return [f"读取文档时出错: {str(e)}"]
    
    # 读取文档
    file_path = r"d:\Ai编程学习\三角函数计算器\多功能计算器软件需求文档（SRD）.docx"
    content = read_word_document(file_path)
    
    print("=== 多功能计算器软件需求文档内容 ===")
    print(f"共读取到 {len(content)} 段内容")
    print("\n")
    
    # 打印前50段内容作为预览
    for i, text in enumerate(content[:50], 1):
        print(f"{i}. {text}")
    
    if len(content) > 50:
        print(f"\n... 还有 {len(content) - 50} 段内容未显示")
    
    # 保存完整内容到文本文件
    with open("document_content.txt", "w", encoding="utf-8") as f:
        for i, text in enumerate(content, 1):
            f.write(f"{i}. {text}\n")
    
    print(f"\n完整内容已保存到 document_content.txt")

except ImportError:
    print("请先安装 python-docx 库: pip install python-docx")
    print("正在尝试使用备用方法...")
    
    # 备用方法：尝试直接提取文本
    import zipfile
    import xml.etree.ElementTree as ET
    
    def extract_text_from_docx(file_path):
        """从docx文件中提取文本"""
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # 读取document.xml
                xml_content = zip_file.read('word/document.xml')
                
                # 解析XML
                root = ET.fromstring(xml_content)
                
                # 定义命名空间
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                # 提取所有文本
                texts = []
                for elem in root.iter():
                    if elem.tag.endswith('}t'):  # 文本元素
                        if elem.text:
                            texts.append(elem.text)
                
                return texts
                
        except Exception as e:
            return [f"提取文本时出错: {str(e)}"]
    
    file_path = r"d:\Ai编程学习\三角函数计算器\多功能计算器软件需求文档（SRD）.docx"
    content = extract_text_from_docx(file_path)
    
    print("=== 多功能计算器软件需求文档内容（备用方法提取） ===")
    print(f"共提取到 {len(content)} 段文本")
    print("\n")
    
    # 打印内容
    for i, text in enumerate(content[:100], 1):
        print(f"{i}. {text}")
    
    if len(content) > 100:
        print(f"\n... 还有 {len(content) - 100} 段内容未显示")